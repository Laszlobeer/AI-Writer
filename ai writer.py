import sys
import requests
from datetime import datetime
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QLabel, QPushButton, QTextEdit,
                             QComboBox, QMessageBox, QFrame, QSlider, QFileDialog,
                             QStatusBar, QSplitter, QSizePolicy, QSpinBox, 
                             QGroupBox, QProgressBar, QShortcut, QScrollArea,
                             QCheckBox, QTabWidget)
from PyQt5.QtCore import QThread, pyqtSignal, Qt, QSize, QTimer
from PyQt5.QtGui import QFont, QTextCursor, QKeySequence

# Try to import docx for .doc support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- Configuration ---
OLLAMA_URL = "http://localhost:11434"
DEFAULT_TOKEN_LIMIT = 140
DEFAULT_TEMPERATURE = 0.7
DEFAULT_CONTEXT_CHARS = 2000  # How many previous characters to include in context

# --- Improved Genre Instructions ---
GENRE_INSTRUCTIONS = {
    "Neutral": """You are an expert ghostwriter. Your task is to continue the provided text seamlessly.
ANALYZE: Study the tone, vocabulary, sentence structure, and pacing of the existing text.
MATCH: Mimic the writing style exactly. Do not shift to a more formal or casual tone.
CONTINUE: Pick up exactly where the text ends. Do not repeat the last sentence.
CONSTRAINTS: Output RAW TEXT ONLY. No introductions, no summaries, no moralizing, and no meta-commentary.""",
    "Dramatic": """You are a dramatic novelist. Continue the text with high emotional intensity.
FOCUS: Conflict, strong feelings, interpersonal tension, and stakes.
STYLE: Use evocative language, internal monologue, and charged dialogue.
CONSTRAINTS: Output RAW TEXT ONLY. Do not resolve the conflict immediately; maintain the tension. No meta-commentary.""",
    "Action": """You are an action thriller writer. Continue the text with fast-paced energy.
FOCUS: Physical movement, kinetics, adrenaline, and immediate danger.
STYLE: Use strong verbs, short punchy sentences, and sensory details (sound, impact).
CONSTRAINTS: Output RAW TEXT ONLY. Keep the pace moving. No long introspections. No meta-commentary.""",
    "Horror": """You are a horror author. Continue the text building dread and atmosphere.
FOCUS: Fear, the unknown, psychological tension, and unsettling imagery.
STYLE: Slow pacing, descriptive shadows, sounds, and feelings of unease.
CONSTRAINTS: Output RAW TEXT ONLY. Do not reveal the monster/threat too quickly. No meta-commentary.""",
    "Romantic": """You are a romance novelist. Continue the text focusing on relationships and emotion.
FOCUS: Chemical connection, intimacy, longing, and emotional vulnerability.
STYLE: Warm, sensory, and character-driven language.
CONSTRAINTS: Output RAW TEXT ONLY. Focus on the dynamic between characters. No meta-commentary.""",
    "Thriller": """You are a thriller writer. Continue the text with suspense and high stakes.
FOCUS: Plot twists, urgency, danger, and clever maneuvering.
STYLE: Tight pacing, cliffhangers, and limited perspective.
CONSTRAINTS: Output RAW TEXT ONLY. Keep the reader guessing. No meta-commentary.""",
    "Comedy": """You are a comedy writer. Continue the text with humor and wit.
FOCUS: Timing, irony, absurdity, or character quirks.
STYLE: Lighthearted, playful, and engaging.
CONSTRAINTS: Output RAW TEXT ONLY. Do not explain the jokes. No meta-commentary.""",
    "SciFi": """You are a science fiction author. Continue the text with futuristic or technological themes.
FOCUS: Science, technology, space, or speculative society elements.
STYLE: Precise terminology, world-building consistency, and logical extrapolation.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain established lore. No meta-commentary.""",
    "Fantasy": """You are a fantasy author. Continue the text with magical or mythical elements.
FOCUS: Magic systems, mythical creatures, quests, or supernatural forces.
STYLE: Epic, descriptive, and immersive language.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain established magic rules. No meta-commentary.""",
    "Crime": """You are a crime fiction writer. Continue the text focusing on criminal activity or investigation.
FOCUS: Detectives, criminals, motives, evidence, or underworld dynamics.
STYLE: Gritty, realistic, or procedural depending on the text's tone.
CONSTRAINTS: Output RAW TEXT ONLY. No meta-commentary.""",
    "Mystery": """You are a mystery writer. Continue the text by introducing or developing clues.
FOCUS: Enigmas, puzzles, hidden information, and deduction.
STYLE: Careful pacing, red herrings, and observational details.
CONSTRAINTS: Output RAW TEXT ONLY. Do not solve the mystery immediately. No meta-commentary.""",
    "Suspense": """You are a suspense writer. Continue the text building anxiety and anticipation.
FOCUS: Ticking clocks, impending doom, and uncertainty.
STYLE: Withholding information, focusing on character worry, and environmental tension.
CONSTRAINTS: Output RAW TEXT ONLY. No meta-commentary.""",
    "Afrofuturism": """You are an Afrofuturism writer. Continue the text blending African diaspora culture with technology and futurism.
FOCUS: African aesthetics, mythology, technology, liberation, and cultural identity.
STYLE: Rich cultural references, vibrant imagery, and speculative elements rooted in African traditions.
CONSTRAINTS: Output RAW TEXT ONLY. Honor cultural authenticity. No meta-commentary.""",
    "Steampunk": """You are a steampunk author. Continue the text with Victorian-era aesthetics and steam-powered technology.
FOCUS: Gears, brass, steam engines, Victorian fashion, and industrial revolution elements.
STYLE: Ornate descriptions, period-appropriate language, and imaginative machinery.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain historical-futuristic blend. No meta-commentary.""",
    "Cyberpunk": """You are a cyberpunk writer. Continue the text with high-tech, low-life themes.
FOCUS: Corporate dominance, cybernetics, hackers, neon-lit cities, and dystopian society.
STYLE: Gritty, noir-influenced, tech-heavy vocabulary, and urban atmosphere.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain the dark futuristic tone. No meta-commentary.""",
    "Futuristic": """You are a futuristic fiction writer. Continue the text with advanced technology and future society elements.
FOCUS: Space travel, AI, advanced weapons, future politics, and human evolution.
STYLE: Clean, forward-thinking language with speculative technology descriptions.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain consistency with established future setting. No meta-commentary.""",
    "Modern": """You are a contemporary fiction writer. Continue the text with realistic, present-day settings and themes.
FOCUS: Current social issues, everyday life, modern relationships, and realistic scenarios.
STYLE: Natural dialogue, relatable situations, and current cultural references.
CONSTRAINTS: Output RAW TEXT ONLY. Keep it grounded in reality. No meta-commentary.""",
    "Historical Fiction": """You are a historical fiction author. Continue the text with accurate period details and settings.
FOCUS: Historical accuracy, period-appropriate language, customs, and social norms.
STYLE: Immersive world-building that reflects the chosen time period authentically.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain historical consistency. No anachronisms. No meta-commentary.""",
    "Philosophy": """You are a philosophical writer. Continue the text exploring deep questions about existence, ethics, and meaning.
FOCUS: Abstract concepts, moral dilemmas, consciousness, truth, and human nature.
STYLE: Thoughtful, reflective, and intellectually engaging prose.
CONSTRAINTS: Output RAW TEXT ONLY. Weave philosophy into narrative naturally. No meta-commentary.""",
    "Dieselpunk": """You are a dieselpunk author. Continue the text with interwar-era aesthetics and diesel-powered technology.
FOCUS: 1920s-1950s aesthetics, heavy machinery, art deco, war machines, and industrial power.
STYLE: Gritty, mechanical descriptions, period slang, and noir influences.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain the diesel-age atmosphere. No meta-commentary.""",
    "Biopunk": """You are a biopunk writer. Continue the text with biotechnology and genetic engineering themes.
FOCUS: DNA manipulation, bio-engineering, corporate biotech, mutations, and organic technology.
STYLE: Clinical yet visceral descriptions, scientific terminology, and body horror elements.
CONSTRAINTS: Output RAW TEXT ONLY. Maintain the bio-tech dystopian tone. No meta-commentary."""
}

# --- Modern Stylesheets ---
STYLES = {
    "light": """
    QMainWindow { background-color: #f9f9fb; }
    QWidget { font-family: 'Segoe UI', 'Roboto', 'Helvetica', sans-serif; color: #2d2d2d; }
    
    QWidget#header { 
        background-color: #ffffff; 
        border-bottom: 1px solid #e0e0e0; 
        padding: 10px; 
    }
    QPushButton { 
        background-color: #ffffff; 
        border: 1px solid #d1d1d1; 
        border-radius: 6px; 
        padding: 6px 12px; 
        font-weight: 500; 
        color: #333; 
    }
    QPushButton:hover { background-color: #f0f0f0; border-color: #bbb; }
    QPushButton:pressed { background-color: #e5e5e5; }
    QPushButton:disabled { color: #aaa; border-color: #eee; background-color: #f9f9f9; }
    
    QPushButton#generate-btn { 
        background-color: #007aff; 
        color: white; 
        border: 1px solid #007aff; 
        font-weight: 600; 
        padding: 6px 20px;
    }
    QPushButton#generate-btn:hover { background-color: #0066d6; }
    QPushButton#generate-btn:disabled { background-color: #a0cfff; border-color: #a0cfff; }

    QPushButton#theme-btn { background-color: transparent; border: none; font-size: 18px; padding: 4px; }
    QPushButton#theme-btn:hover { background-color: #f0f0f0; }
    
    QPushButton#memory-btn { 
        background-color: #ff9500; 
        color: white; 
        border: 1px solid #ff9500; 
        font-weight: 600; 
        padding: 6px 15px;
    }
    QPushButton#memory-btn:hover { background-color: #e68600; }

    QTextEdit#editor { 
        background-color: #ffffff; 
        border: none; 
        padding: 30px 40px; 
        font-size: 18px; 
        line-height: 1.6; 
        font-family: 'Georgia', 'Times New Roman', serif; 
        color: #1a1a1a;
        selection-background-color: #b3d7ff;
    }
    QTextEdit#editor:focus { outline: none; }
    QTextEdit#editor:disabled { background-color: #fafafa; color: #999; }
    
    QTextEdit#memory-view { 
        background-color: #f5f5f7; 
        border: 1px solid #e0e0e0; 
        border-radius: 6px; 
        padding: 15px; 
        font-size: 13px; 
        line-height: 1.5; 
        font-family: 'Consolas', 'Monaco', monospace; 
        color: #555;
    }

    QFrame#sidebar { 
        background-color: #ffffff; 
        border-left: 1px solid #e0e0e0; 
        padding: 0px; 
    }
    
    QGroupBox { 
        font-weight: 600; 
        color: #555; 
        border: 1px solid #e0e0e0; 
        border-radius: 8px; 
        margin-top: 12px; 
        padding-top: 10px; 
    }
    QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 5px; }

    QComboBox, QSpinBox { 
        background-color: #f5f5f7; 
        border: 1px solid #d1d1d1; 
        border-radius: 6px; 
        padding: 6px; 
        font-size: 13px; 
        color: #333;
    }
    QComboBox:focus, QSpinBox:focus { border: 1px solid #007aff; background-color: #fff; }
    QComboBox::drop-down { border: none; padding-right: 10px; }
    QComboBox::down-arrow { image: none; border-left: 1px solid #d1d1d1; padding-left: 5px; }

    QSlider::groove:horizontal { border: 1px solid #ddd; height: 6px; background: #e0e0e0; border-radius: 3px; }
    QSlider::handle:horizontal { background: #007aff; border: 1px solid #0056b3; width: 16px; margin: -5px 0; border-radius: 8px; }
    QSlider::handle:horizontal:hover { background: #0056b3; }

    QLabel#sidebar-title { font-size: 13px; font-weight: 600; color: #555; margin: 4px 0; }
    QLabel#status-label { color: #666; font-size: 12px; }
    QLabel#memory-indicator { 
        background-color: #34c759; 
        color: white; 
        border-radius: 10px; 
        padding: 3px 10px; 
        font-size: 11px; 
        font-weight: 600; 
    }
    
    QStatusBar { background-color: #ffffff; border-top: 1px solid #e0e0e0; color: #666; padding: 4px; }
    
    QSplitter::handle { background-color: #e0e0e0; width: 4px; }
    QSplitter::handle:hover { background-color: #007aff; }
    
    QProgressBar { 
        border: 1px solid #ddd; 
        border-radius: 4px; 
        text-align: center; 
        background: #f0f0f0; 
        height: 10px; 
    }
    QProgressBar::chunk { background-color: #007aff; border-radius: 3px; }
    
    QTabWidget::pane { border: 1px solid #e0e0e0; border-radius: 6px; }
    QTabBar::tab { 
        background-color: #f0f0f0; 
        border: 1px solid #e0e0e0; 
        border-radius: 4px 4px 0 0; 
        padding: 6px 12px; 
        margin-right: 2px; 
    }
    QTabBar::tab:selected { background-color: #ffffff; border-bottom: 1px solid #ffffff; }
    QTabBar::tab:hover { background-color: #e5e5e5; }
    """,
    "dark": """
    QMainWindow { background-color: #1e1e1e; }
    QWidget { font-family: 'Segoe UI', 'Roboto', 'Helvetica', sans-serif; color: #d4d4d4; }
    
    QWidget#header { 
        background-color: #252526; 
        border-bottom: 1px solid #3e3e42; 
        padding: 10px; 
    }
    QPushButton { 
        background-color: #3c3c3c; 
        border: 1px solid #3c3c3c; 
        border-radius: 6px; 
        padding: 6px 12px; 
        font-weight: 500; 
        color: #d4d4d4; 
    }
    QPushButton:hover { background-color: #505050; }
    QPushButton:pressed { background-color: #606060; }
    QPushButton:disabled { color: #666; background-color: #2d2d2d; }
    
    QPushButton#generate-btn { 
        background-color: #0e639c; 
        color: white; 
        border: 1px solid #0e639c; 
        font-weight: 600; 
        padding: 6px 20px;
    }
    QPushButton#generate-btn:hover { background-color: #1177bb; }
    QPushButton#generate-btn:disabled { background-color: #3c3c3c; border-color: #3c3c3c; color: #888; }

    QPushButton#theme-btn { background-color: transparent; border: none; font-size: 18px; padding: 4px; }
    QPushButton#theme-btn:hover { background-color: #3c3c3c; }
    
    QPushButton#memory-btn { 
        background-color: #d97706; 
        color: white; 
        border: 1px solid #d97706; 
        font-weight: 600; 
        padding: 6px 15px;
    }
    QPushButton#memory-btn:hover { background-color: #b45309; }

    QTextEdit#editor { 
        background-color: #1e1e1e; 
        border: none; 
        padding: 30px 40px; 
        font-size: 18px; 
        line-height: 1.6; 
        font-family: 'Georgia', 'Times New Roman', serif; 
        color: #d4d4d4;
        selection-background-color: #264f78;
    }
    QTextEdit#editor:focus { outline: none; }
    QTextEdit#editor:disabled { background-color: #1a1a1a; color: #666; }
    
    QTextEdit#memory-view { 
        background-color: #2d2d30; 
        border: 1px solid #3e3e42; 
        border-radius: 6px; 
        padding: 15px; 
        font-size: 13px; 
        line-height: 1.5; 
        font-family: 'Consolas', 'Monaco', monospace; 
        color: #aaa;
    }

    QFrame#sidebar { 
        background-color: #252526; 
        border-left: 1px solid #3e3e42; 
        padding: 0px; 
    }
    
    QGroupBox { 
        font-weight: 600; 
        color: #aaa; 
        border: 1px solid #3e3e42; 
        border-radius: 8px; 
        margin-top: 12px; 
        padding-top: 10px; 
    }
    QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 5px; }

    QComboBox, QSpinBox { 
        background-color: #3c3c3c; 
        border: 1px solid #3c3c3c; 
        border-radius: 6px; 
        padding: 6px; 
        font-size: 13px; 
        color: #d4d4d4;
    }
    QComboBox:focus, QSpinBox:focus { border: 1px solid #0e639c; background-color: #444; }
    QComboBox::drop-down { border: none; padding-right: 10px; }
    QComboBox::down-arrow { image: none; border-left: 1px solid #555; padding-left: 5px; }

    QSlider::groove:horizontal { border: 1px solid #555; height: 6px; background: #3c3c3c; border-radius: 3px; }
    QSlider::handle:horizontal { background: #0e639c; border: 1px solid #0e639c; width: 16px; margin: -5px 0; border-radius: 8px; }
    QSlider::handle:horizontal:hover { background: #1177bb; }

    QLabel#sidebar-title { font-size: 13px; font-weight: 600; color: #aaa; margin: 4px 0; }
    QLabel#status-label { color: #888; font-size: 12px; }
    QLabel#memory-indicator { 
        background-color: #10b981; 
        color: white; 
        border-radius: 10px; 
        padding: 3px 10px; 
        font-size: 11px; 
        font-weight: 600; 
    }
    
    QStatusBar { background-color: #252526; border-top: 1px solid #3e3e42; color: #888; padding: 4px; }
    
    QSplitter::handle { background-color: #3e3e42; width: 4px; }
    QSplitter::handle:hover { background-color: #0e639c; }
    
    QProgressBar { 
        border: 1px solid #555; 
        border-radius: 4px; 
        text-align: center; 
        background: #3c3c3c; 
        height: 10px; 
    }
    QProgressBar::chunk { background-color: #0e639c; border-radius: 3px; }
    
    QTabWidget::pane { border: 1px solid #3e3e42; border-radius: 6px; }
    QTabBar::tab { 
        background-color: #3c3c3c; 
        border: 1px solid #3e3e42; 
        border-radius: 4px 4px 0 0; 
        padding: 6px 12px; 
        margin-right: 2px; 
    }
    QTabBar::tab:selected { background-color: #252526; border-bottom: 1px solid #252526; }
    QTabBar::tab:hover { background-color: #505050; }
    """
}

class OllamaWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    models_loaded = pyqtSignal(list)
    progress = pyqtSignal(int)

    def __init__(self, endpoint, model=None, prompt=None, context=None, temperature=0.7, 
                 token_limit=140, genre="Neutral", memory_summary=None):
        super().__init__()
        self.endpoint = endpoint
        self.model = model
        self.prompt = prompt
        self.context = context
        self.temperature = temperature
        self.token_limit = token_limit
        self.genre = genre
        self.memory_summary = memory_summary

    def run(self):
        try:
            if self.endpoint == "scan":
                response = requests.get(f"{OLLAMA_URL}/api/tags", timeout=5)
                if response.status_code == 200:
                    data = response.json()
                    models = [m['name'] for m in data.get('models', [])]
                    self.models_loaded.emit(models)
                else:
                    self.error.emit(f"API Error: {response.status_code}")
            
            elif self.endpoint == "generate":
                system_instruction = GENRE_INSTRUCTIONS.get(self.genre, GENRE_INSTRUCTIONS["Neutral"])
                
                # Build prompt with memory context
                prompt_parts = [system_instruction]
                
                if self.memory_summary:
                    prompt_parts.append(f"\n\n### STORY SUMMARY (Memory) ###\n{self.memory_summary}\n")
                
                if self.context:
                    prompt_parts.append(f"\n\n### PREVIOUS CONTEXT ###\n{self.context}\n")
                
                prompt_parts.append(f"\n\n### CURRENT TEXT ###\n{self.prompt}\n")
                prompt_parts.append(f"\n### CONTINUATION ###\n")
                
                full_prompt = "".join(prompt_parts)
            
                payload = {
                    "model": self.model,
                    "prompt": full_prompt,
                    "stream": False,
                    "options": {
                        "num_predict": self.token_limit,
                        "temperature": self.temperature
                    }
                }
                response = requests.post(f"{OLLAMA_URL}/api/generate", json=payload, timeout=120)
                if response.status_code == 200:
                    data = response.json()
                    completion = data.get('response', '')
                    cleaned = self.clean_completion(self.prompt, completion)
                    self.finished.emit(cleaned)
                else:
                    self.error.emit(f"Generation Error: {response.status_code}")
        except requests.exceptions.ConnectionError:
            self.error.emit("Cannot connect to Ollama. Is it running?")
        except Exception as e:
            self.error.emit(str(e))

    def clean_completion(self, original, completion):
        original_stripped = original.strip()
        completion_stripped = completion.strip()

        if completion_stripped.startswith(original_stripped):
            completion_stripped = completion_stripped[len(original_stripped):].strip()

        prefixes = [
            "here's the continuation: ", "continuation: ", "continued: ", 
            "here is the completion: ", "here's the completion: ",
            "the continuation is: ", "completion: ", "### Continuation ###",
            "### CONTINUATION ###"
        ]
        for prefix in prefixes:
            if completion_stripped.lower().startswith(prefix):
                completion_stripped = completion_stripped[len(prefix):].strip()

        if completion_stripped.startswith('"') and not original_stripped.endswith('"'):
            completion_stripped = completion_stripped[1:].strip()

        return completion_stripped

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Writer Pro - With Memory")
        self.setMinimumSize(1200, 800)
        self.current_theme = "light"
        self.temperature = DEFAULT_TEMPERATURE
        self.token_limit = DEFAULT_TOKEN_LIMIT
        self.context_chars = DEFAULT_CONTEXT_CHARS
        self.current_file = None
        self.selected_genre = "Neutral"
        self.memory_enabled = True
        self.generation_history = []  # Track all generations in session
        self.setStyleSheet(STYLES[self.current_theme])
        
        self.init_ui()
        self.scan_models()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # --- Custom Header Bar ---
        self.header = QWidget()
        self.header.setObjectName("header")
        self.header.setFixedHeight(60)
        header_layout = QHBoxLayout(self.header)
        header_layout.setContentsMargins(15, 10, 15, 10)
        
        # Theme Toggle
        self.theme_btn = QPushButton("🌙")
        self.theme_btn.setObjectName("theme-btn")
        self.theme_btn.setFixedWidth(40)
        self.theme_btn.clicked.connect(self.toggle_theme)
        header_layout.addWidget(self.theme_btn)
        
        header_layout.addSpacing(10)
        
        # Model Selection
        self.model_combo = QComboBox()
        self.model_combo.setMinimumWidth(200)
        self.model_combo.addItem("Select model...")
        header_layout.addWidget(QLabel("  Model:"))
        header_layout.addWidget(self.model_combo)
        
        self.scan_btn = QPushButton("🔄")
        self.scan_btn.setToolTip("Refresh Models")
        self.scan_btn.setFixedWidth(40)
        self.scan_btn.clicked.connect(self.scan_models)
        header_layout.addWidget(self.scan_btn)
        
        header_layout.addSpacing(20)
        
        # Genre Dropdown
        self.genre_combo = QComboBox()
        self.genre_combo.setMinimumWidth(180)
        self.genre_combo.addItems(list(GENRE_INSTRUCTIONS.keys()))
        self.genre_combo.setCurrentText("Neutral")
        self.genre_combo.currentTextChanged.connect(self.on_genre_changed)
        header_layout.addWidget(QLabel("  📚 Genre:"))
        header_layout.addWidget(self.genre_combo)
        
        header_layout.addSpacing(20)
        
        # Memory Toggle
        self.memory_btn = QPushButton("🧠 Memory: ON")
        self.memory_btn.setObjectName("memory-btn")
        self.memory_btn.setCheckable(True)
        self.memory_btn.setChecked(True)
        self.memory_btn.clicked.connect(self.toggle_memory)
        header_layout.addWidget(self.memory_btn)
        
        header_layout.addStretch()
        
        # Generate Button
        self.generate_btn = QPushButton("✨ Generate")
        self.generate_btn.setObjectName("generate-btn")
        self.generate_btn.setFixedHeight(35)
        self.generate_btn.clicked.connect(self.start_generation)
        self.generate_btn.setEnabled(False)
        header_layout.addWidget(self.generate_btn)
        
        header_layout.addSpacing(10)
        
        # Save Buttons
        self.save_txt_btn = QPushButton("📄 .txt")
        self.save_txt_btn.clicked.connect(self.save_txt)
        header_layout.addWidget(self.save_txt_btn)
        
        if DOCX_AVAILABLE:
            self.save_doc_btn = QPushButton("📕 .docx")
            self.save_doc_btn.clicked.connect(self.save_docx)
            header_layout.addWidget(self.save_doc_btn)
        
        main_layout.addWidget(self.header)

        # --- Main Content Splitter ---
        splitter = QSplitter(Qt.Horizontal)
        splitter.setHandleWidth(6)
        
        # Editor Area with Memory Tab
        editor_container = QWidget()
        editor_layout = QVBoxLayout(editor_container)
        editor_layout.setSpacing(0)
        editor_layout.setContentsMargins(0, 0, 0, 0)
        
        # Tab Widget for Editor and Memory View
        self.editor_tabs = QTabWidget()
        
        # Main Editor Tab
        self.editor = QTextEdit()
        self.editor.setObjectName("editor")
        self.editor.setPlaceholderText("Start writing your story here...\n\nPress Ctrl+Enter to let AI continue from your cursor position.\n\n🧠 Memory is ENABLED - AI will remember previous content!")
        self.editor.textChanged.connect(self.on_text_changed)
        self.editor_tabs.addTab(self.editor, "📝 Editor")
        
        # Memory Context Tab
        self.memory_view = QTextEdit()
        self.memory_view.setObjectName("memory-view")
        self.memory_view.setReadOnly(True)
        self.memory_view.setPlaceholderText("Memory context will appear here when you generate...\n\nThis shows what the AI can see from your previous writing.")
        self.editor_tabs.addTab(self.memory_view, "🧠 Memory Context")
        
        # Generation History Tab
        self.history_view = QTextEdit()
        self.history_view.setObjectName("memory-view")
        self.history_view.setReadOnly(True)
        self.history_view.setPlaceholderText("Generation history will appear here...")
        self.editor_tabs.addTab(self.history_view, "📜 History")
        
        editor_layout.addWidget(self.editor_tabs)
        splitter.addWidget(editor_container)
        
        # Sidebar
        sidebar = QFrame()
        sidebar.setObjectName("sidebar")
        sidebar.setFixedWidth(320)
        sidebar_layout = QVBoxLayout(sidebar)
        sidebar_layout.setContentsMargins(15, 15, 15, 15)
        sidebar_layout.setSpacing(15)
        
        # Memory Settings Group
        memory_group = QGroupBox("🧠 Memory Settings")
        memory_layout = QVBoxLayout(memory_group)
        
        # Memory Status Indicator
        self.memory_status = QLabel("✅ Active")
        self.memory_status.setObjectName("memory-indicator")
        self.memory_status.setAlignment(Qt.AlignCenter)
        memory_layout.addWidget(self.memory_status)
        
        # Context Length Control
        context_label = QLabel("📊 Context Length (Characters)")
        context_label.setObjectName("sidebar-title")
        memory_layout.addWidget(context_label)
        
        self.context_spinbox = QSpinBox()
        self.context_spinbox.setMinimum(100)
        self.context_spinbox.setMaximum(10000)
        self.context_spinbox.setValue(DEFAULT_CONTEXT_CHARS)
        self.context_spinbox.setSuffix(" chars")
        self.context_spinbox.valueChanged.connect(self.on_context_changed)
        memory_layout.addWidget(self.context_spinbox)
        
        context_hint = QLabel("How much previous text AI can see")
        context_hint.setObjectName("status-label")
        memory_layout.addWidget(context_hint)
        
        # Clear Memory Button
        self.clear_memory_btn = QPushButton("🗑️ Clear Memory")
        self.clear_memory_btn.clicked.connect(self.clear_memory)
        memory_layout.addWidget(self.clear_memory_btn)
        
        sidebar_layout.addWidget(memory_group)
        
        # Temperature Group
        temp_group = QGroupBox("🌡️ Creativity (Temperature)")
        temp_layout = QVBoxLayout(temp_group)
        
        self.temp_value_label = QLabel(f"{self.temperature:.2f}")
        self.temp_value_label.setAlignment(Qt.AlignRight)
        self.temp_value_label.setObjectName("sidebar-title")
        temp_layout.addWidget(self.temp_value_label)
        
        self.temp_slider = QSlider(Qt.Horizontal)
        self.temp_slider.setMinimum(0)
        self.temp_slider.setMaximum(200)
        self.temp_slider.setValue(int(DEFAULT_TEMPERATURE * 100))
        self.temp_slider.setTickPosition(QSlider.TicksBelow)
        self.temp_slider.setTickInterval(25)
        self.temp_slider.valueChanged.connect(self.on_temperature_changed)
        temp_layout.addWidget(self.temp_slider)
        
        temp_hints = QHBoxLayout()
        temp_hints.addWidget(QLabel("🎯 Focused"))
        temp_hints.addStretch()
        temp_hints.addWidget(QLabel("🎨 Creative"))
        temp_layout.addLayout(temp_hints)
        
        sidebar_layout.addWidget(temp_group)
        
        # Token Limit Group
        token_group = QGroupBox("📊 Length (Tokens)")
        token_layout = QVBoxLayout(token_group)
        
        self.token_spinbox = QSpinBox()
        self.token_spinbox.setMinimum(10)
        self.token_spinbox.setMaximum(2000)
        self.token_spinbox.setValue(DEFAULT_TOKEN_LIMIT)
        self.token_spinbox.setSuffix(" tokens")
        self.token_spinbox.valueChanged.connect(self.on_token_limit_changed)
        token_layout.addWidget(self.token_spinbox)
        
        token_hints = QLabel("Higher = longer response\nLower = shorter response")
        token_hints.setObjectName("status-label")
        token_layout.addWidget(token_hints)
        
        sidebar_layout.addWidget(token_group)
        
        # Tips Group
        tips_group = QGroupBox("ℹ️ Tips")
        tips_layout = QVBoxLayout(tips_group)
        tips = [
            "• Memory helps AI maintain continuity",
            "• Place cursor where AI should continue",
            "• Ctrl+Enter to generate",
            "• Lower temp = consistent style",
            "• Higher temp = surprising ideas",
            "• Check Memory tab to see context",
            "• Clear memory for new chapters"
        ]
        for tip in tips:
            lbl = QLabel(tip)
            lbl.setObjectName("status-label")
            lbl.setWordWrap(True)
            tips_layout.addWidget(lbl)
        
        sidebar_layout.addWidget(tips_group)
        sidebar_layout.addStretch()
        
        splitter.addWidget(sidebar)
        main_layout.addWidget(splitter)
        
        # Set initial splitter sizes (70% editor, 30% sidebar)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 0)
        
        # Status Bar
        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.char_label = QLabel("0 chars | 0 words")
        self.char_label.setObjectName("status-label")
        self.statusBar.addPermanentWidget(self.char_label)
        self.statusBar.showMessage("🧠 Memory System Ready")
        
        # Keyboard Shortcuts
        shortcut = QShortcut(QKeySequence("Ctrl+Return"), self)
        shortcut.activated.connect(self.start_generation)
        
        # Progress Indicator (Hidden by default)
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedWidth(150)
        self.progress_bar.setRange(0, 0)
        self.progress_bar.hide()
        self.statusBar.addPermanentWidget(self.progress_bar)
        
        # Update memory view initially
        self.update_memory_view()

    def on_genre_changed(self, genre):
        self.selected_genre = genre
        self.statusBar.showMessage(f"📚 Genre set to: {genre}")

    def toggle_theme(self):
        if self.current_theme == "light":
            self.current_theme = "dark"
            self.theme_btn.setText("☀️")
        else:
            self.current_theme = "light"
            self.theme_btn.setText("🌙")
        self.setStyleSheet(STYLES[self.current_theme])

    def toggle_memory(self):
        self.memory_enabled = not self.memory_enabled
        if self.memory_enabled:
            self.memory_btn.setText("🧠 Memory: ON")
            self.memory_status.setText("✅ Active")
            self.statusBar.showMessage("🧠 Memory enabled - AI will remember context")
        else:
            self.memory_btn.setText("🧠 Memory: OFF")
            self.memory_status.setText("⏸️ Paused")
            self.statusBar.showMessage("⏸️ Memory disabled - AI won't see previous context")
        self.update_memory_view()

    def on_context_changed(self, value):
        self.context_chars = value
        self.statusBar.showMessage(f"🧠 Context length set to {self.context_chars} characters")
        self.update_memory_view()

    def on_temperature_changed(self, value):
        self.temperature = value / 100.0
        self.temp_value_label.setText(f"{self.temperature:.2f}")

    def on_token_limit_changed(self, value):
        self.token_limit = value
        self.statusBar.showMessage(f"Token limit set to {self.token_limit} tokens")

    def on_text_changed(self):
        text = self.editor.toPlainText()
        char_count = len(text)
        word_count = len(text.split())
        self.char_label.setText(f"{char_count} chars | {word_count} words")
        
        model = self.model_combo.currentText()
        has_text = len(text.strip()) > 0
        has_model = model not in ["Select model...", "No models found"]
        self.generate_btn.setEnabled(has_text and has_model and not self.progress_bar.isVisible())
        
        # Update memory view when text changes
        if self.memory_enabled:
            self.update_memory_view()

    def update_memory_view(self):
        """Update the memory context view to show what AI will see"""
        text = self.editor.toPlainText()
        
        if not self.memory_enabled:
            self.memory_view.setText("⏸️ Memory is currently disabled.\n\nEnable it to let AI remember your previous writing.")
            return
        
        if not text.strip():
            self.memory_view.setText("📝 Start writing to build memory context...\n\nThe AI will use recent text to maintain continuity.")
            return
        
        # Show the context that will be sent to AI
        context_start = max(0, len(text) - self.context_chars)
        context_text = text[context_start:]
        
        memory_info = f"""🧠 MEMORY CONTEXT PREVIEW
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Context Length: {len(context_text)} / {self.context_chars} characters
Total Document: {len(text)} characters
Generations in Session: {len(self.generation_history)}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{context_text if context_text else "(No context yet)"}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💡 The AI will see this context + your current cursor position to continue writing seamlessly.
"""
        self.memory_view.setText(memory_info)

    def clear_memory(self):
        """Clear the generation history (not the document)"""
        reply = QMessageBox.question(self, 'Clear Memory', 
                                    'Clear generation history? This will reset what the AI remembers from this session.\n\nYour document will NOT be deleted.',
                                    QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.generation_history = []
            self.history_view.clear()
            self.update_memory_view()
            self.statusBar.showMessage("🗑️ Memory cleared - Starting fresh context")
            QMessageBox.information(self, "Memory Cleared", "Generation history has been cleared.\n\nYour document is safe - only the AI's session memory was reset.")

    def scan_models(self):
        self.statusBar.showMessage("🔍 Scanning for models...")
        self.model_combo.clear()
        self.model_combo.addItem("Scanning...")
        self.generate_btn.setEnabled(False)
        
        self.worker = OllamaWorker(endpoint="scan")
        self.worker.models_loaded.connect(self.on_models_loaded)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_models_loaded(self, models):
        self.model_combo.clear()
        if not models:
            self.model_combo.addItem("No models found")
            self.statusBar.showMessage("❌ No models available")
            self.generate_btn.setEnabled(False)
        else:
            self.model_combo.addItems(models)
            self.statusBar.showMessage(f"✓ {len(models)} models available | 🧠 Memory Ready")
            self.on_text_changed()

    def start_generation(self):
        text = self.editor.toPlainText()
        model = self.model_combo.currentText()
        
        if not text.strip() or model in ["Select model...", "No models found"]:
            QMessageBox.warning(self, "Warning", "Please enter text and select a model.")
            return

        self.generation_cursor_pos = len(text)
        
        # Get context for memory
        context = None
        if self.memory_enabled and len(text) > self.context_chars:
            context_start = max(0, len(text) - self.context_chars)
            context = text[context_start:self.generation_cursor_pos]
        elif self.memory_enabled:
            context = text[:self.generation_cursor_pos]
        
        # Update memory view with actual context being sent
        if self.memory_enabled and context:
            self.memory_view.setText(f"📤 SENDING TO AI:\n\n{context[-500:] if len(context) > 500 else context}\n\n...")
        
        # UI State: Loading
        self.generate_btn.setEnabled(False)
        self.generate_btn.setText("⏳ Writing...")
        self.progress_bar.show()
        memory_status = "🧠 With Memory" if self.memory_enabled else "⏸️ No Memory"
        self.statusBar.showMessage(f"AI is writing ({memory_status}, {self.selected_genre}, Temp: {self.temperature:.2f})...")
        self.editor.setEnabled(False)
        
        self.worker = OllamaWorker(endpoint="generate", model=model, prompt=text[self.generation_cursor_pos:], 
                                   context=context, temperature=self.temperature, token_limit=self.token_limit, 
                                   genre=self.selected_genre)
        self.worker.finished.connect(self.on_generation_finished)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_generation_finished(self, completion):
        # UI State: Ready
        self.progress_bar.hide()
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("✨ Generate")
        self.editor.setEnabled(True)
        self.editor.setFocus()
        
        if not completion.strip():
            self.statusBar.showMessage("⚠️ No completion generated")
            return
        
        # Add to generation history
        self.generation_history.append({
            'timestamp': datetime.now().strftime("%H:%M:%S"),
            'length': len(completion),
            'genre': self.selected_genre
        })
        
        # Update history tab
        history_text = "📜 GENERATION HISTORY\n" + "="*50 + "\n\n"
        for i, gen in enumerate(self.generation_history[-10:], 1):  # Show last 10
            history_text += f"{i}. [{gen['timestamp']}] {gen['length']} chars - {gen['genre']}\n"
        self.history_view.setText(history_text)
        
        # Insert completion at cursor position
        cursor = self.editor.textCursor()
        cursor.setPosition(self.generation_cursor_pos)
        
        text = self.editor.toPlainText()
        if text and not text[-1].isspace() and completion and not completion[0].isspace():
            cursor.insertText(" ")
        
        cursor.insertText(completion)
        self.editor.setTextCursor(cursor)
        
        self.statusBar.showMessage(f"✓ Completion added ({len(completion)} chars) | 🧠 Memory Updated")
        self.on_text_changed()
        self.update_memory_view()

    def on_error(self, error_msg):
        self.progress_bar.hide()
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("✨ Generate")
        self.editor.setEnabled(True)
        self.statusBar.showMessage("❌ Error")
        QMessageBox.critical(self, "Error", error_msg)

    def save_txt(self):
        text = self.editor.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "Warning", "No text to save!")
            return
        
        default_name = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save as TXT", default_name, "Text Files (*.txt)")
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                self.statusBar.showMessage(f"✓ Saved to {file_path}")
                self.current_file = file_path
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save: {str(e)}")

    def save_docx(self):
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Warning", "python-docx not installed. Run: pip install python-docx")
            return
        
        text = self.editor.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "Warning", "No text to save!")
            return
        
        default_name = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save as DOCX", default_name, "Word Documents (*.docx)")
        
        if file_path:
            try:
                doc = Document()
                doc.add_heading('AI Writer Document', 0)
                doc.add_paragraph(f'Created: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
                doc.add_paragraph()
                
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
                
                doc.save(file_path)
                self.statusBar.showMessage(f"✓ Saved to {file_path}")
                self.current_file = file_path
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save: {str(e)}")

if __name__ == "__main__":
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
    
    app = QApplication(sys.argv)
    font = QFont("Segoe UI", 10)
    app.setFont(font)
    
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())